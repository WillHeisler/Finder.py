#This project was started out of frustration of downloading
#files, but not remembering where the file was saved and/or
#what the file was named. This issue made it nearly impossible
#to find these files. Thus finder was born.

#Through various use cases, the project morphed from
#a program to only search for files to one that can
#be used to do dirty word searches across file shares
#or a local computer. While the original functionality
#still exists, the command line argument options have
#expanded the possible uses.

import re
import pathlib
import os
import argparse
import sys
import docx
"""
Next addition to this project is
to add csv/xlsx support for searching
files.
"""

def _init_menu():
    """
    initialize the command line parameter menu.
    """
    menu_parser = argparse.ArgumentParser()
    menu_parser.add_argument("-f", "--file", metavar='',
                             help = "Optional. File name to search for. Enter 'unknown' if unknown.")
    menu_parser.add_argument("-s", "--string", metavar='',
                             help = """Not used if -iL is used. The string to search for.
                             Enter 'none' if not desired.""")
    menu_parser.add_argument("-p", "--path", metavar='', required=True,
                             help = "Required. Directory path to search in.")
    menu_parser.add_argument("-o", "--output", metavar='', required=True,
                             help = "Required. Absolute file path to save output")
    menu_parser.add_argument("-iL", "--inputList", metavar='' ,
                             help = """Not used if -s is used. Output file to dump output.
                             Include extension.""")
    return menu_parser
def _menu(menu_parser):
    """
    Parser for the command line parameter menu and calls the appropriate functions.
    :param menu_parser: the argparse menu as created with '_init_menu()'
    """
    args = menu_parser.parse_args()
    if args.file:
        question_file_name = args.file
    else:
        question_file_name = 'unknown'
    if args.string:
        string_search = args.string
    if args.path:
        question_file_path = args.path
    if args.output:
        output_file = pathlib.Path(args.output)
        if not output_file.exists():
            output_file.touch()
    if args.inputList:
        fh = pathlib.Path(args.inputList)
        fc = fh.read_text()
        string_search = fc.split("\n")
    return question_file_path, question_file_name, string_search, str(output_file)
def question_answers(file_path, file_name, stng):
    """
    function that emulates Linux utilities 'grep' and 'find'
    sets fn and stng to 'none' in case 'Enter' is pressed instead of
    entering 'none'.
    """
    print("\n\n\n\n\n")
    file_path = pathlib.Path(file_path)
    file_name = file_name.lower()
    if not file_name.strip():
        file_name = 'none'
    try:
        if not stng.strip():
            stng = 'none'
        stng = stng.lower()
    except:
        pass
    return file_path, file_name, stng
def list_dir(file_path, file_name):
    """
    checks if no filename was given and searches for the directory
    if directory exists and is found, prints all contents recursively
    """
    try:
        if file_name == 'none':
            if not file_path.is_dir():
                print(f"The directory {file_path} does not exist.")
            for entry in file_path.iterdir():
                if entry.is_file():
                    print(f"File: {entry}")
                elif entry.is_dir():
                    print(f"Directory: {entry}")
    except Exception as exception:
        print("The search for the directory faile. It may be a permissions issue")
        print(exception)
def find_file(file_path, file_name, output_file):
    """
    checks if filename was given and searches for the file
    if file is not found, prints it is not a valid filename
    """
    output_file = pathlib.Path(output_file)
    try:
        if file_name not in ['none', 'unknown']:
            file_path_name = file_path / file_name
            if file_path_name.is_file():
                full_path = file_path_name
                print("I found the file!!!!!!!")
                print(str(file_path_name) + "\n")
                return full_path
            if file_path.is_dir():
                file_path_list = []
                for root, dirs, files in os.walk(file_path):
                    for file in files:
                        if file_name in file:
                            full_path = pathlib.Path(os.path.join(root, file))
                            file_path_list.append(full_path)
                            print("I found the file!!!!!!")
                        else:
                            file_pattern = f'(.*?{file_name}.*?)'
                            if re.findall(file_pattern, file, flags=re.IGNORECASE):
                                full_path = pathlib.Path(os.path.join(root, file))
                                file_path_list.append(full_path)
                                print("I found the file!!!!!!")
                file_path_string = "\n".split(file_path_list)
                if output_file:
                    output_file.write_text(file_path_string)
                else:
                    print(file_path_string)
            else:
                print(f"{file_name} is not a valid filename.")
    except Exception as exception:
        print("There has been an error. This might be a permissions error")
        print(exception)
def grep(file_name, stng, output_file):
    """
    checks if there was a string given to search
    if the string is found, prints the lines that contain the string
    if it is not found, prints that it can't be found or doesn't exist
    """
    output_file = pathlib.Path(output_file)
    if stng not in ['none', 'unkown']:
        ext = pathlib.Path(file_name).suffix.lower()
        pattern = f'(.*?{stng}.*?)'
        if ext in ['.txt', '.docx']:
            if ext == '.txt':
                content = pathlib.Path(file_name).read_text()
                search = re.findall(pattern, content, flags=re.IGNORECASE)
            elif ext == '.docx':
                doc = docx.Document(file_name).paragraphs
                search = []
                for paragraph in doc:
                    text_search = (re.findall(pattern, paragraph.text, flags=re.IGNORECASE))
                    if len(text_search) != 0:
                        search.append(text_search)
        if len(search) != 0:
            results_list = []
            for lines in search:
                results_list.append(str(lines))
            clean_results = "\n".join(results_list)
            if output_file:
                output_file.write_text(clean_results)
                print("Output saved in the output file specified.")
            else:
                print(f"I found '{stng}' {len(search)} times!")
                print("Here are all the lines that is was found on.\n")
                print(clean_results.strip())
        else:
            print(f"The string {stng} was not found")
def mystery_file(file_path, stng, output_file):
    """Looking for a unknown filename inside parent
    dir using a known string within the file
    """
    output_file = pathlib.Path(output_file)
    parent_path = pathlib.Path(file_path)
    if parent_path.is_dir():
        mystery_file_list = []
        for root, dirs, files in os.walk(parent_path):
            for file in files:
                path = os.path.join(root, file)
                ext = pathlib.Path(file).suffix.lower()
                strng_pattern = f'(.*?{stng}.*?)'
                if ext in ['.txt', '.docx']:
                    if ext == '.txt':
                        try:
                            fh = pathlib.Path(root) / file
                            fc = fh.read_text()
                            lines = fc.split('\n')
                            doc_list = []
                            for line in lines:
                                if not fh in doc_list:
                                    if re.match(strng_pattern, line, flags=re.IGNORECASE):
                                        doc_list.append(fh)
                                        mystery_file_list.append(str(fh))
                        except:
                            pass
                    elif ext == '.docx':
                        try:
                            doc = docx.Document(path)
                            list_of_docs = []
                            if doc not in list_of_docs:
                                for paragraph in doc.paragraphs:
                                    if re.search(rf'(.*{stng}.*)', paragraph.text, flags=re.IGNORECASE):
                                        list_of_docs.append(file)
                                        mystery_file_list.append(path)
                        except:
                            pass
        mystery_file_string = "\n".join(list(set(mystery_file_list)))
        if output_file.exists():
            output_file.write_text(mystery_file_string)
        else:
            print(mystery_file_string)
if __name__ == '__main__':
    """
    questions for the user on what directory, filename, and/or string to search for
    """
    if len(sys.argv) == 1:
        print("""This program is able to list the contents of a directory,
search for a file by name recursively from a known directory,
search for a string in a known file and directory, or find an
unknown file based on a known string when given a parent directory.\n\n""")
        question_file_name = input("""What is the name of the file that you would like to find?
If you are not looking for a specific file enter 'None', if
name is unkown, enter 'Unknown'.\n""")
        print("""\n\nWhat is the file path you would like for me to search?
Be as specific as possible.""", end="")
        question_file_path = input("""The more broad the search, the 
longer and more likely for falses.
ie 'C:\\Users\\%USER%\\Downloads'.\n""")
        string_search = input("\n\nWhat is the string you want to search for? default is 'None'. "
"This option does not support lists.\n")
        output_file =input("\n\nWhat is the output directory and filename you want findings savings?\n")
    else:
        question_file_path, question_file_name, string_search, output_file = _menu(_init_menu())
    file_path, file_name, stng = question_answers(question_file_path, question_file_name, string_search)
    if question_file_path and not question_file_name and not stng:
        list_dir(file_path, file_name)
    elif question_file_name and question_file_path and (string_search == 'none'): #or string_search != string_search.strip()):
        find_file(file_path, file_name, output_file)
    elif question_file_name.lower() == 'unknown':
        mystery_file(file_path, stng, output_file)
    elif question_file_path and question_file_name and string_search != 'none':
        grep(find_file(file_path, file_name, output_file), stng, output_file)
        